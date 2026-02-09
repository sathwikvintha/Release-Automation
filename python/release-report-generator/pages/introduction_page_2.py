from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns


RED = "FF0000"
GREY = "D9D9D9"
BLACK = RGBColor(0, 0, 0)


# ==================================================
# Reuse SAME helper functions (copy-paste identical)
# ==================================================
def _colored_bar(doc, text, bg_color, bold=True, italic=False, size=11, center=False, add_borders=False):
    """
    Full-width colored bar (used for section headers)
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), bg_color)
    pPr.append(shd)

    if add_borders:
        pBdr = OxmlElement("w:pBdr")

        top = OxmlElement("w:top")
        top.set(ns.qn("w:val"), "single")
        top.set(ns.qn("w:sz"), "12")
        top.set(ns.qn("w:space"), "1")
        top.set(ns.qn("w:color"), "000000")
        pBdr.append(top)

        bottom = OxmlElement("w:bottom")
        bottom.set(ns.qn("w:val"), "single")
        bottom.set(ns.qn("w:sz"), "12")
        bottom.set(ns.qn("w:space"), "1")
        bottom.set(ns.qn("w:color"), "000000")
        pBdr.append(bottom)

        pPr.append(pBdr)


def _body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)

    run = p.runs[0]
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK


# ==================================================
# INTRODUCTION – PAGE 2
# ==================================================
def build_introduction_page_2(doc):
    """
    Builds:
    1.7 CSC ORM
    1.8 NewUX
    1.9 CSCApp
    1.10 CSC FO Scheduler
    """

    # ============================
    # 1.7 CSC ORM
    # ============================
    _colored_bar(doc, "1.7 CSC ORM", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for CSC ORM deployment attached:"
    )

    attach = doc.add_paragraph("• Connexis Supply Chain Front Office")
    attach.paragraph_format.left_indent = Pt(30)
    attach.runs[0].font.size = Pt(10.5)

    # ============================
    # 1.8 NewUX
    # ============================
    _colored_bar(doc, "1.8 NewUX", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for NewUX deployment attached:\n\n"
        "Yet to be release for R26.1"
    )

    # ============================
    # 1.9 CSCApp
    # ============================
    _colored_bar(doc, "1.9 CSCApp", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for CSCApp deployment attached:\n\n"
        "Yet to be release for R26.1"
    )

    # ============================
    # 1.10 CSC FO Scheduler
    # ============================
    _colored_bar(doc, "1.10 CSC FO Scheduler", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for CSC FO Scheduler deployment attached:\n\n"
        "Yet to be release for R26.1"
    )

    # End of Introduction section
    doc.add_page_break()

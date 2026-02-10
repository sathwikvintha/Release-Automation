from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns


RED = "FF0000"
GREY = "D9D9D9"
BLACK = RGBColor(0, 0, 0)


def _colored_bar(doc, text, bg_color, bold=True, italic=False, size=11, center=False, add_borders=False):
    """
    Full-width colored bar (used for section headers)
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Center alignment if requested
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), bg_color)
    pPr.append(shd)
    
    # Add borders if requested
    if add_borders:
        pBdr = OxmlElement("w:pBdr")
        
        # Top border
        top = OxmlElement("w:top")
        top.set(ns.qn("w:val"), "single")
        top.set(ns.qn("w:sz"), "12")  # Border width
        top.set(ns.qn("w:space"), "1")
        top.set(ns.qn("w:color"), "000000")  # Black border
        pBdr.append(top)
        
        # Bottom border
        bottom = OxmlElement("w:bottom")
        bottom.set(ns.qn("w:val"), "single")
        bottom.set(ns.qn("w:sz"), "12")  # Border width
        bottom.set(ns.qn("w:space"), "1")
        bottom.set(ns.qn("w:color"), "000000")  # Black border
        pBdr.append(bottom)
        
        pPr.append(pBdr)


def _body_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(2)

    run = p.runs[0]
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK


def build_introduction_page_1(doc):
    """
    Builds:
    1. Introduction
    1.1 Purpose
    1.2 CSC INT API
    1.3 Marketing Web Page
    1.4 Supplier Onboarding
    1.5 CSC CR
    1.6 CSC FO
    """

    # ============================
    # 1. INTRODUCTION (RED BAR) - CENTERED WITH BORDERS
    # ============================
    _colored_bar(doc, "1.  Introduction", RED, bold=True, size=13, center=True, add_borders=True)

    # ============================
    # 1.1 Purpose
    # ============================
    _colored_bar(doc, "1.1 Purpose", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "The objective of this document is to articulate the deployment instructions for FO 26.1 release.\n"
        "We will follow same steps as it was provided in 25.3 deployment document and this document covers\n"
        "specific steps for 26.1 release deployment information."
    )

    # ============================
    # 1.2 CSC INT API
    # ============================
    _colored_bar(doc, "1.2 CSC INT API", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for CSC INT API deployment attached:\n\n"
        "Yet to be release for R26.1"
    )

    # ============================
    # 1.3 Marketing WEB PAGE
    # ============================
    _colored_bar(doc, "1.3 Marketing WEB PAGE", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for Marketing webpage deployment attached:\n\n"
        "Yet to be release for R26.1"
    )

    # ============================
    # 1.4 Supplier Onboarding
    # ============================
    _colored_bar(doc, "1.4 Supplier Onboarding", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for Supplier Onboarding deployment attached:"
    )

    # Placeholder for attachment (editable by user)
    attach = doc.add_paragraph("• Supplier_Onboarding – Deployment_Doc")
    attach.paragraph_format.left_indent = Pt(30)
    attach.runs[0].font.size = Pt(10.5)

    # ============================
    # 1.5 CSC CR
    # ============================
    _colored_bar(doc, "1.5 CSC CR", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for CSC CR deployment attached:"
    )

    attach2 = doc.add_paragraph("• Connexis Supply Chain Client Registration")
    attach2.paragraph_format.left_indent = Pt(30)
    attach2.runs[0].font.size = Pt(10.5)

    # ============================
    # 1.6 CSC FO
    # ============================
    _colored_bar(doc, "1.6 CSC FO", GREY, italic=True, size=11)

    _body_paragraph(
        doc,
        "Please find the detailed information document for CSC FO deployment attached:\n\n"
        "Yet to be release for R26.1"
    )

    # Page break for Introduction – Page 2
    doc.add_page_break()
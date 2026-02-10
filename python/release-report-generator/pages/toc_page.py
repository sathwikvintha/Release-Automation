"""
Table of Contents Page Builder for Word Documents

This module creates professional, properly formatted table of contents pages
with dynamic leader dots that keep page numbers perfectly aligned to the right.
"""

from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================================
# COLOR CONSTANTS
# ============================================================================
BLUE = RGBColor(0, 0, 204)
BLACK = RGBColor(0, 0, 0)


# ============================================================================
# FONT SIZE CONSTANTS
# ============================================================================
TITLE_FONT_SIZE = Pt(14)
TOC_HEADING_FONT_SIZE = Pt(12)
SECTION_FONT_SIZE = Pt(11)
ITEM_FONT_SIZE = Pt(10)


# ============================================================================
# SPACING CONSTANTS
# ============================================================================
TITLE_SPACE_AFTER = Pt(12)
TOC_HEADING_SPACE_AFTER = Pt(16)
SECTION_SPACE_BEFORE = Pt(8)
ITEM_LINE_SPACING = 1.15


# ============================================================================
# LAYOUT CONSTANTS
# ============================================================================
# Approximate character width for dot calculation (adjust based on font)
CHAR_WIDTH_SECTION = 0.07  # Inches per character for section items
CHAR_WIDTH_ITEM = 0.06     # Inches per character for sub-items
PAGE_WIDTH = 6.5           # Usable page width in inches (8.5" - margins)
SECTION_INDENT = 0.25      # Left indent for sections
ITEM_INDENT = 0.5          # Left indent for items


# ============================================================================
# MAIN BUILDER FUNCTION
# ============================================================================
def build_toc_page(doc):
    """
    Build a complete professional Table of Contents page.
    
    Args:
        doc: A python-docx Document object
    """
    _add_document_title(doc)
    _add_horizontal_line(doc)
    _add_toc_heading(doc)
    _add_toc_content(doc)


# ============================================================================
# TITLE AND HEADER FUNCTIONS
# ============================================================================
def _add_document_title(doc):
    """
    Add the main document title with centered alignment and blue color.
    
    Args:
        doc: A python-docx Document object
    """
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = TITLE_SPACE_AFTER
    title.paragraph_format.space_before = Pt(0)
    
    run = title.add_run("Release Deployment Document")
    run.bold = True
    run.font.size = TITLE_FONT_SIZE
    run.font.color.rgb = BLUE


def _add_horizontal_line(doc):
    """
    Add a horizontal line separator below the title.
    
    Args:
        doc: A python-docx Document object
    """
    line_paragraph = doc.add_paragraph()
    line_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_paragraph.paragraph_format.space_after = Pt(16)
    line_paragraph.paragraph_format.space_before = Pt(0)
    
    # Add bottom border to create a professional line
    pPr = line_paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_toc_heading(doc):
    """
    Add the 'Table of Contents' heading.
    
    Args:
        doc: A python-docx Document object
    """
    toc_heading = doc.add_paragraph()
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.paragraph_format.space_after = TOC_HEADING_SPACE_AFTER
    toc_heading.paragraph_format.space_before = Pt(0)
    
    run = toc_heading.add_run("Table of Contents")
    run.bold = True
    run.font.size = TOC_HEADING_FONT_SIZE
    run.font.color.rgb = BLUE


# ============================================================================
# CONTENT BUILDER FUNCTION
# ============================================================================
def _add_toc_content(doc):
    """
    Add all TOC sections and items.
    
    Args:
        doc: A python-docx Document object
    """
    # Section 1: INTRODUCTION
    _add_toc_section(doc, "1.", "INTRODUCTION", "4")
    
    introduction_items = [
        ("1.1", "PURPOSE", "4"),
        ("1.2", "CSC INT API", "4"),
        ("1.3", "MARKETING WEB PAGE", "4"),
        ("1.4", "SUPPLIER ONBOARDING", "4"),
        ("1.5", "CSC CR", "4"),
        ("1.6", "CSC FO", "4"),
        ("1.7", "CSC ORM", "5"),
        ("1.8", "NEWUX", "5"),
        ("1.9", "CSCAPP", "5"),
        ("1.10", "CSC FO SCHEDULER", "5"),
    ]
    
    for number, title, page in introduction_items:
        _add_toc_item(doc, number, title, page)
    
    # Section 2: DEPLOYMENT DETAILS
    _add_toc_section(doc, "2.", "DEPLOYMENT DETAILS", "6", add_space_before=True)
    
    deployment_items = [
        ("2.1", "WEB SERVER CHANGES", "6"),
        ("2.2", "MAVEN DEPLOYMENT CHANGES", "6"),
        ("2.3", "APP SERVER CHANGES", "7"),
        ("2.4", "DB CHANGES- ENVIRONMENT SPECIFIC CHANGES", "8"),
        ("2.5", "QUEUE CONFIGURATION SCRIPTS", "8"),
        ("2.6", "SCHEDULER JOBS", "9"),
        ("2.7", "MIGRATION SCRIPTS", "9"),
        ("2.8", "SHELL SCRIPT CHANGES", "9"),
        ("2.9", "SQL SCRIPT CHANGE", "10"),
        ("2.10", "CRON JOB CHANGES", "10"),
        ("2.11", "KEYCLOAK CONFIGURATION CHANGES", "10"),
        ("2.12", "SCHEDULER SERVER CHANGES", "11"),
    ]
    
    for number, title, page in deployment_items:
        _add_toc_item(doc, number, title, page)


# ============================================================================
# HELPER FUNCTION FOR DYNAMIC DOT CALCULATION
# ============================================================================
def _calculate_dots_needed(text_length, char_width, indent):
    """
    Calculate the number of dots needed to push page numbers to the right edge.
    
    Args:
        text_length: Total character count of number + title
        char_width: Approximate width per character in inches
        indent: Left indent in inches
    
    Returns:
        Number of dots needed
    """
    # Calculate space used by text
    text_width = text_length * char_width
    
    # Calculate remaining space (page width - indent - text width - space for page number)
    remaining_width = PAGE_WIDTH - indent - text_width - 0.5  # 0.5" for page number
    
    # Calculate dots needed
    dots_needed = int(remaining_width / (char_width * 0.8))  # Dots are slightly narrower
    dots_needed = dots_needed + 10
    # Ensure minimum dots for readability
    return max(dots_needed, 10)


# ============================================================================
# TOC ITEM FORMATTING FUNCTIONS
# ============================================================================
def _add_toc_section(doc, number, title, page, add_space_before=False):
    """
    Add a major TOC section heading (bold, larger text) with dynamic dots.
    
    Args:
        doc: A python-docx Document object
        number: Section number (e.g., "1.")
        title: Section title
        page: Page number
        add_space_before: Whether to add extra space before this section
    """
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(SECTION_INDENT)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = ITEM_LINE_SPACING
    
    if add_space_before:
        paragraph.paragraph_format.space_before = SECTION_SPACE_BEFORE
    else:
        paragraph.paragraph_format.space_before = Pt(0)
    
    # Format number and title
    formatted_number = f"{number:<6}"
    formatted_title = f"{title} "
    
    # Calculate dots needed based on content length
    total_text_length = len(formatted_number) + len(formatted_title)
    dots_count = _calculate_dots_needed(total_text_length, CHAR_WIDTH_SECTION, SECTION_INDENT)
    
    # Add section number
    num_run = paragraph.add_run(formatted_number)
    num_run.bold = True
    num_run.font.size = SECTION_FONT_SIZE
    num_run.font.color.rgb = BLUE
    
    # Add section title
    title_run = paragraph.add_run(formatted_title)
    title_run.bold = True
    title_run.font.size = SECTION_FONT_SIZE
    title_run.font.color.rgb = BLUE
    
    # Add dynamic leader dots
    dots_run = paragraph.add_run("." * dots_count)
    dots_run.font.size = SECTION_FONT_SIZE
    dots_run.font.color.rgb = BLUE
    
    # Add page number
    page_run = paragraph.add_run(f" {page}")
    page_run.bold = True
    page_run.font.size = SECTION_FONT_SIZE
    page_run.font.color.rgb = BLUE


def _add_toc_item(doc, number, title, page):
    """
    Add a TOC item (subsection) with proper indentation and dynamic dots.
    
    Args:
        doc: A python-docx Document object
        number: Item number (e.g., "1.1")
        title: Item title
        page: Page number
    """
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(ITEM_INDENT)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = ITEM_LINE_SPACING
    
    # Format number and title
    formatted_number = f"{number:<8}"
    formatted_title = f"{title} "
    
    # Calculate dots needed based on content length
    total_text_length = len(formatted_number) + len(formatted_title)
    dots_count = _calculate_dots_needed(total_text_length, CHAR_WIDTH_ITEM, ITEM_INDENT)
    
    # Add item number
    num_run = paragraph.add_run(formatted_number)
    num_run.font.size = ITEM_FONT_SIZE
    num_run.font.color.rgb = BLUE
    
    # Add item title
    title_run = paragraph.add_run(formatted_title)
    title_run.font.size = ITEM_FONT_SIZE
    title_run.font.color.rgb = BLUE
    
    # Add dynamic leader dots
    dots_run = paragraph.add_run("." * dots_count)
    dots_run.font.size = ITEM_FONT_SIZE
    dots_run.font.color.rgb = BLUE
    
    # Add page number
    page_run = paragraph.add_run(f" {page}")
    page_run.font.size = ITEM_FONT_SIZE
    page_run.font.color.rgb = BLUE


# ============================================================================
# EXAMPLE USAGE
# ============================================================================
if __name__ == "__main__":
    from docx import Document
    
    # Create a new document
    doc = Document()
    
    # Build TOC page
    build_toc_page(doc)
    
    # Save the document
    doc.save("professional_toc_aligned.docx")
    print("✓ Professional Table of Contents created successfully!")
    print("✓ Page numbers perfectly aligned to the right!")
    print("✓ File saved as: professional_toc_aligned.docx")
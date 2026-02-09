from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement, ns

GREY = "D9D9D9"
BLACK = RGBColor(0, 0, 0)


# --------------------------------------------------
# Helpers
# --------------------------------------------------
def _colored_bar(doc, text, bg_color, italic=True, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)

    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), bg_color)
    pPr.append(shd)


def _release_block(doc, release, text, first=False):
    """
    Release number + tight underline + aligned instruction
    """

    space_before = Pt(6) if first else Pt(2)

    # ----------------------------
    # Release label
    # ----------------------------
    p_rel = doc.add_paragraph()
    p_rel.paragraph_format.left_indent = Pt(24)
    p_rel.paragraph_format.space_before = space_before
    p_rel.paragraph_format.space_after = Pt(2)

    run = p_rel.add_run(release)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    run.underline = True

    # ----------------------------
    # Instruction
    # ----------------------------
    p_txt = doc.add_paragraph()
    p_txt.paragraph_format.left_indent = Pt(42)
    p_txt.paragraph_format.space_before = Pt(1)
    p_txt.paragraph_format.space_after = Pt(8)

    content = text.strip() if text else "No special instructions"
    r_txt = p_txt.add_run(content)
    r_txt.font.size = Pt(10.5)
    r_txt.font.color.rgb = BLACK


def _get_release_text(data, section, release):
    section_data = data.get("deploymentDetails", {}).get(section)

    if isinstance(section_data, dict):
        return section_data.get(release, "")

    if isinstance(section_data, list):
        return "\n".join(section_data) if section_data else ""

    return ""


# --------------------------------------------------
# PAGE BUILDER
# --------------------------------------------------
def build_deployment_details_page_3(doc, data):
    """
    Deployment Details – Page 3

    NOTE:
    - For 2.5 Queue Configuration Scripts,
      R26.1.0.5 is intentionally SKIPPED
    """

    releases = [
        "R26.1.0.1",
        "R26.1.0.2",
        "R26.1.0.3",
        "R26.1.0.4",
        "R26.1.0.5",
    ]

    # ==================================================
    # RELEASE NOTES FROM 2.3 APP SERVER CHANGES
    # ==================================================
    for idx, rel in enumerate(releases):
        text = _get_release_text(
            data,
            "2.3 App Server Changes",
            rel
        )
        _release_block(doc, rel, text, first=(idx == 0))

    # ==================================================
    # 2.4 DB Changes – Environment Specific Changes
    # ==================================================
    _colored_bar(
        doc,
        "2.4 DB Changes- Environment Specific Changes",
        GREY,
        italic=True
    )

    for idx, rel in enumerate(releases):
        text = _get_release_text(
            data,
            "2.4 DB Changes- Environment Specific Changes",
            rel
        )
        _release_block(doc, rel, text, first=(idx == 0))

    # ==================================================
    # 2.5 Queue Configuration Scripts
    # (INTENTIONALLY SKIP R26.1.0.5)
    # ==================================================
    _colored_bar(
        doc,
        "2.5 Queue Configuration Scripts",
        GREY,
        italic=True
    )

    # Only up to R26.1.0.4
    for idx, rel in enumerate(releases[:-1]):
        text = _get_release_text(
            data,
            "2.5 Queue Configuration Scripts",
            rel
        )
        _release_block(doc, rel, text, first=(idx == 0))


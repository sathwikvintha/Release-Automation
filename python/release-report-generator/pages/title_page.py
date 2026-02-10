from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def build_title_page(doc, title, release, subtitle, version_no, version_date):
    doc.add_paragraph("\n\n")

    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.bold = True
    r.font.size = Pt(22)

    doc.add_paragraph("\n")

    p = doc.add_paragraph(f"Release {release}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    doc.add_paragraph("\n")

    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph(
        f"Version No : {version_no}\nVersion Date : {version_date}"
    )
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.bold = True
        run.font.size = Pt(11)

    doc.add_page_break()

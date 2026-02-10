from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, ns

def build_document_control_page(doc, version_no):
    doc.add_heading("Document Control", level=2)

    table = doc.add_table(3, 3)
    table.style = "Table Grid"

    for row in table.rows:
        row.height = Inches(0.8)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    labels = [
        ("Author(s):", "", ""),
        ("Creation Date:", "Updated by:", "Reviewed by:"),
        (f"Revision No: {version_no}", "Updated On:", "Reviewed On:")
    ]

    for r, row in enumerate(labels):
        for c, val in enumerate(row):
            run = table.cell(r, c).paragraphs[0].add_run(val)
            run.bold = c == 0
            run.font.size = Pt(11)

    doc.add_paragraph("\n")
    doc.add_heading("Revision History", level=2)

    rh = doc.add_table(1, 4)
    rh.style = "Table Grid"

    # Set column widths for better spacing
    rh.columns[0].width = Inches(1.2)  # Date
    rh.columns[1].width = Inches(1.0)  # Version
    rh.columns[2].width = Inches(3.5)  # Description
    rh.columns[3].width = Inches(1.5)  # Author

    # Style headers
    headers = ["Date", "Version", "Description", "Author"]
    for i, h in enumerate(headers):
        cell = rh.rows[0].cells[i]
        cell.text = h
        
        # Set row height for header
        rh.rows[0].height = Inches(0.4)
        
        # Center align and bold header text
        paragraph = cell.paragraphs[0]
        paragraph.alignment = 1  # Center alignment
        run = paragraph.runs[0]
        run.bold = True
        run.font.size = Pt(11)
        
        # Add gray background
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(ns.qn("w:fill"), "D9D9D9")
        tcPr.append(shd)
        
        # Vertical center alignment
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Add 5 data rows with proper spacing
    for _ in range(5):
        row = rh.add_row()
        row.height = Inches(0.5)  # Set consistent row height
        
        # Apply vertical alignment and padding to each cell
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Add some padding to paragraphs
            paragraph = cell.paragraphs[0]
            paragraph_format = paragraph.paragraph_format
            paragraph_format.space_before = Pt(6)
            paragraph_format.space_after = Pt(6)
            paragraph_format.line_spacing = 1.15

    doc.add_page_break()
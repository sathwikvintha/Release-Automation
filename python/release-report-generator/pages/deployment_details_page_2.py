from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, ns

# ==================================================
# COLORS
# ==================================================
GREY = "D9D9D9"
YELLOW = "FFFF00"
BLUE = RGBColor(0, 0, 255)
BLACK = RGBColor(0, 0, 0)

# ==================================================
# INTERNAL HELPERS
# ==================================================
def _colored_bar(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), GREY)
    pPr.append(shd)


def _shade(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), color)
    tcPr.append(shd)


def _border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(ns.qn("w:val"), "single")
        b.set(ns.qn("w:sz"), "4")
        b.set(ns.qn("w:color"), "000000")
        borders.append(b)
    tcPr.append(borders)


def _format(cell, text, *, bold=False, underline=False, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.runs[0]
    r.font.size = Pt(10.5)
    r.font.color.rgb = color
    r.bold = bold
    r.underline = underline


def _set_tbl_grid(table, widths):
    tbl = table._tbl
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(ns.qn("w:w"), str(int(w.twips)))
        tblGrid.append(gridCol)
    tbl.append(tblGrid)


def _set_row_height(row, height_twips=430):
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(ns.qn("w:val"), str(height_twips))
    trHeight.set(ns.qn("w:hRule"), "exact")
    trPr.append(trHeight)

# ==================================================
# PAGE BUILDER
# ==================================================
def build_deployment_details_page_2(doc, data):

    # --------------------------------------------------
    # SECTION HEADER
    # --------------------------------------------------
    _colored_bar(doc, "2.3 App Server Changes")

    intro = doc.add_paragraph("Please find the application build version details:")
    intro.runs[0].bold = True
    intro.runs[0].font.size = Pt(10.5)

    # --------------------------------------------------
    # TABLE 1 – MODULE VS VERSION
    # --------------------------------------------------
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = "Table Grid"
    table1.allow_autofit = False

    _set_tbl_grid(table1, [Inches(3.2), Inches(1.8)])

    headers1 = ["Application Module", "Latest Release Build version"]
    for i, h in enumerate(headers1):
        c = table1.rows[0].cells[i]
        _format(c, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(c, GREY)
        _border(c)

    modules = [
        "OLDUX MyBatis Hibernate",
        "NEWUX-Angular",
        "CR Hibernate",
        "ORM",
        "API",
        "SUPPLIER ONBOARDING",
        "CSC App Hibernate",
        "Marketing Web Page",
        "CSC INT API",
        "Keycloak",
    ]

    versions = data.get("appServerBuildVersions", {})

    for m in modules:
        r = table1.add_row()
        _set_row_height(r)

        _format(r.cells[0], m, underline=True, color=BLUE)
        _border(r.cells[0])

        v = versions.get(m, "")
        _format(r.cells[1], v, color=BLUE)
        _border(r.cells[1])

        if v == "R26.1.0.5":
            _shade(r.cells[1], YELLOW)

    doc.add_paragraph("")

    # --------------------------------------------------
    # TABLE 2 – RELEASE WISE MODULES (Sr No FIXED)
    # --------------------------------------------------
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    table2.allow_autofit = False

    # 🔥 ULTRA-NARROW Sr No + EXTRA SPACE TO MODULES
    _set_tbl_grid(
        table2,
        [
            Inches(0.18),  # Sr No (tight)
            Inches(1.9),   # Release Build No
            Inches(3.4),   # Application Modules (wide)
        ]
    )

    headers2 = ["Sr No.", "Release Build No.", "Application Modules"]
    for i, h in enumerate(headers2):
        c = table2.rows[0].cells[i]
        _format(c, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade(c, GREY)
        _border(c)

    releases = data.get("releaseWiseModules", [])

    if not releases:
        releases = [
            {"release": "Release FO R26.1.0.1", "modules": ["CR - clientreg-ear.ear, DB Script"]},
            {"release": "Release FO R26.1.0.2", "modules": ["CSCSO - cscsoapp.war"]},
            {"release": "Release FO R26.1.0.3", "modules": [
                "OldUX – bnppng.war, db Scripts",
                "MDP – bnp_middleware_mdp.war",
                "Scheduler libs",
                "NewUX – bnppngux-ear.ear",
                "CR – clientreg-ear.ear"
            ]},
            {"release": "Release FO R26.1.0.4", "modules": [
                "ORM – ear, DB Scripts",
                "CR – clientreg-ear.ear"
            ]},
            {"release": "Release FO R26.1.0.5", "modules": [
                "ORM – ear, DB Scripts",
                "CR – clientreg-ear.ear",
                "CSCSO – cscsoapp.war, DB Scripts"
            ]},
        ]

    for idx, rel in enumerate(releases, start=1):
        r = table2.add_row()

        _format(r.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        _border(r.cells[0])

        _format(r.cells[1], rel["release"])
        _border(r.cells[1])

        cell = r.cells[2]
        cell.text = ""
        for m in rel["modules"]:
            p = cell.add_paragraph(f"• {m}")
            p.paragraph_format.left_indent = Inches(0.12)
            p.runs[0].font.size = Pt(10.5)
        _border(cell)


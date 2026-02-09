from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns


RED = "FF0000"
GREY = "D9D9D9"
BLACK = RGBColor(0, 0, 0)

RELEASE_VERSIONS = [
    "R26.1.0.1",
    "R26.1.0.2",
    "R26.1.0.3",
    "R26.1.0.4",
    "R26.1.0.5",
]


# ==================================================
# COMMON HELPERS
# ==================================================
def _colored_bar(doc, text, bg_color, bold=True, italic=False, size=11, center=False, add_borders=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), bg_color)
    pPr.append(shd)

    if add_borders:
        pBdr = OxmlElement("w:pBdr")
        for side in ("top", "bottom"):
            el = OxmlElement(f"w:{side}")
            el.set(ns.qn("w:val"), "single")
            el.set(ns.qn("w:sz"), "12")
            el.set(ns.qn("w:space"), "1")
            el.set(ns.qn("w:color"), "000000")
            pBdr.append(el)
        pPr.append(pBdr)


def _release_block(doc, release, instructions):
    """
    R26.x.x.x  (text-only underline, very close)
    No special instructions
    """

    # ----------------------------
    # Release label (TEXT underline ONLY)
    # ----------------------------
    p_rel = doc.add_paragraph()
    p_rel.paragraph_format.left_indent = Pt(18)
    p_rel.paragraph_format.space_before = Pt(0)
    p_rel.paragraph_format.space_after = Pt(1)

    run = p_rel.add_run(release)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    run.underline = True   # ✅ CORRECT underline

    # ----------------------------
    # Instruction text
    # ----------------------------
    p_txt = doc.add_paragraph(instructions)
    p_txt.paragraph_format.left_indent = Pt(36)
    p_txt.paragraph_format.space_before = Pt(1)
    p_txt.paragraph_format.space_after = Pt(8)

    run_txt = p_txt.runs[0]
    run_txt.font.size = Pt(10.5)
    run_txt.font.color.rgb = BLACK


def _get_instruction_for_release(section_data, release):
    for item in section_data:
        if release in item:
            return item
    return "No special instructions"


# ==================================================
# DEPLOYMENT DETAILS – PAGE 1
# ==================================================
def build_deployment_details_page_1(doc, json_data):

    deployment = json_data.get("deploymentDetails", {})

    # ============================
    # 2. DEPLOYMENT DETAILS
    # ============================
    _colored_bar(
        doc,
        "2.  Deployment Details",
        RED,
        bold=True,
        size=13,
        center=True,
        add_borders=True,
    )

    # ============================
    # 2.1 Web server Changes
    # ============================
    section_name = "2.1 Web server Changes"
    _colored_bar(doc, section_name, GREY, italic=True, size=11)

    section_data = deployment.get(section_name, [])

    for rel in RELEASE_VERSIONS:
        instr = _get_instruction_for_release(section_data, rel)
        _release_block(doc, rel, instr)

    # ============================
    # 2.2 Maven Deployment Changes
    # ============================
    section_name = "2.2 Maven Deployment Changes"
    _colored_bar(doc, section_name, GREY, italic=True, size=11)

    section_data = deployment.get(section_name, [])

    for rel in RELEASE_VERSIONS:
        instr = _get_instruction_for_release(section_data, rel)
        _release_block(doc, rel, instr)

    doc.add_page_break()

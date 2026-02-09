from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement, ns

GREY = "D9D9D9"
BLACK = RGBColor(0, 0, 0)

RELEASES = [
    "R26.1.0.1",
    "R26.1.0.2",
    "R26.1.0.3",
    "R26.1.0.4",
    "R26.1.0.5",
]


# --------------------------------------------------
# Helpers
# --------------------------------------------------
def _colored_bar(doc, text, bg_color, italic=True, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK

    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(ns.qn("w:fill"), bg_color)
    pPr.append(shd)


def _release_block(doc, release, text, first=False):
    """
    R26.x.x.x (tight underline)
    Instruction text
    """
    # Release label
    p_rel = doc.add_paragraph()
    p_rel.paragraph_format.left_indent = Pt(24)
    p_rel.paragraph_format.space_before = Pt(6 if first else 4)
    p_rel.paragraph_format.space_after = Pt(1)

    r = p_rel.add_run(release)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLACK
    r.underline = True

    # Instruction text
    p_txt = doc.add_paragraph()
    p_txt.paragraph_format.left_indent = Pt(42)
    p_txt.paragraph_format.space_before = Pt(0)
    p_txt.paragraph_format.space_after = Pt(6)

    content = text.strip() if text else "No special instructions"
    run_txt = p_txt.add_run(content)
    run_txt.font.size = Pt(10.5)
    run_txt.font.color.rgb = BLACK


def _get_release_text(data, section, release):
    section_data = data.get("deploymentDetails", {}).get(section)

    if isinstance(section_data, dict):
        return section_data.get(release, "")

    if isinstance(section_data, list):
        return "\n".join(section_data) if section_data else ""

    return ""


# --------------------------------------------------
# PAGE BUILDER
# --------------------------------------------------
def build_deployment_details_page_6(doc, data):
    """
    Deployment Details – Page 6

    Includes:
    - 2.11 Keycloak Configuration changes
    - 2.12 Scheduler Server changes
    """

    # ==================================================
    # 2.11 Keycloak Configuration changes
    # ==================================================
    _colored_bar(doc, "2.11 Keycloak Configuration changes", GREY, italic=True)

    for idx, rel in enumerate(RELEASES):
        text = _get_release_text(
            data,
            "2.11 Keycloak Configuration changes",
            rel
        )
        _release_block(doc, rel, text, first=(idx == 0))

    # ==================================================
    # 2.12 Scheduler Server changes
    # ==================================================
    _colored_bar(doc, "2.12 Scheduler Server changes", GREY, italic=True)

    for idx, rel in enumerate(RELEASES):
        text = _get_release_text(
            data,
            "2.12 Scheduler Server changes",
            rel
        )
        _release_block(doc, rel, text, first=(idx == 0))

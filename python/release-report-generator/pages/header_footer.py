from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from config.paths import BNP_LOGO, ORACLE_LOGO, FOOTER_LINE

def add_header(section):
    header = section.header
    p = header.paragraphs[0]
    run = p.add_run("Confidential - Oracle Restricted")
    run.bold = True
    run.font.size = Pt(11)

    table = header.add_table(1, 2, Inches(6))
    table.cell(0, 0).paragraphs[0].add_run().add_picture(BNP_LOGO, width=Inches(3))
    table.cell(0, 1).paragraphs[0].add_run().add_picture(ORACLE_LOGO, width=Inches(2))
    table.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sep = header.add_paragraph()
    pPr = sep._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(ns.qn("w:val"), "single")
    b.set(ns.qn("w:sz"), "6")
    pBdr.append(b)
    pPr.append(pBdr)


def add_footer(section, doc_name):
    footer = section.footer
    usable_width = section.page_width - section.left_margin - section.right_margin

    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run().add_picture(FOOTER_LINE, width=usable_width)

    p = footer.add_paragraph()
    p.paragraph_format.space_after = Pt(0)

    left = p.add_run(doc_name)
    left.font.size = Pt(9)

    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(ns.qn("w:val"), "right")
    tab.set(ns.qn("w:pos"), str(int(usable_width / 635)))
    tabs.append(tab)
    pPr.append(tabs)

    p.add_run("\t")

    r = p.add_run("Page ")
    r.font.size = Pt(9)

    fld1 = OxmlElement("w:fldChar")
    fld1.set(ns.qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar")
    fld2.set(ns.qn("w:fldCharType"), "end")
    r._r.extend([fld1, instr, fld2])
